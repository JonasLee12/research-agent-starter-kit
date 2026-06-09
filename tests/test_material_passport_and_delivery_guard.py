#!/usr/bin/env python3
"""Tests for material passport and formal delivery guard helpers."""

from __future__ import annotations

import subprocess
import sys
import tempfile
import unittest
import uuid
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))
from _docx_runtime import BUNDLED_PYTHON


PYTHON = sys.executable
DOCX_PYTHON = str(BUNDLED_PYTHON if BUNDLED_PYTHON.exists() else Path(sys.executable))


class MaterialPassportAndDeliveryGuardTests(unittest.TestCase):
    def run_script(self, args: list[str]) -> subprocess.CompletedProcess[str]:
        return subprocess.run(args, cwd=ROOT, text=True, capture_output=True)

    def write_basic_docx(self, path: Path, text: str = "Clean draft.") -> None:
        code = """
from docx import Document
import sys
doc = Document()
doc.add_paragraph(sys.argv[2])
doc.save(sys.argv[1])
"""
        proc = subprocess.run([DOCX_PYTHON, "-c", code, str(path), text], text=True, capture_output=True)
        self.assertEqual(proc.returncode, 0, proc.stderr + proc.stdout)

    def write_pipe_table_docx(self, path: Path) -> None:
        code = """
from docx import Document
import sys
doc = Document()
for text in sys.argv[2:]:
    doc.add_paragraph(text)
doc.save(sys.argv[1])
"""
        proc = subprocess.run(
            [DOCX_PYTHON, "-c", code, str(path), "| A | B |", "|---|---|", "| one | two |"],
            text=True,
            capture_output=True,
        )
        self.assertEqual(proc.returncode, 0, proc.stderr + proc.stdout)

    def write_real_table_docx(self, path: Path) -> None:
        code = """
from docx import Document
import sys
doc = Document()
table = doc.add_table(rows=2, cols=2)
table.cell(0, 0).text = "A"
table.cell(0, 1).text = "B"
table.cell(1, 0).text = "one"
table.cell(1, 1).text = "two"
doc.save(sys.argv[1])
"""
        proc = subprocess.run([DOCX_PYTHON, "-c", code, str(path)], text=True, capture_output=True)
        self.assertEqual(proc.returncode, 0, proc.stderr + proc.stdout)

    def write_flat_heading_docx(self, path: Path) -> None:
        code = """
from docx import Document
import sys
doc = Document()
doc.styles["Heading 1"].font.bold = True
heading = doc.add_heading(level=1)
run = heading.add_run("Flat heading")
run.font.bold = False
doc.add_paragraph("Body.")
doc.save(sys.argv[1])
"""
        proc = subprocess.run([DOCX_PYTHON, "-c", code, str(path)], text=True, capture_output=True)
        self.assertEqual(proc.returncode, 0, proc.stderr + proc.stdout)

    def test_short_material_passport_warns_for_missing_artifact(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp) / "passport.md"
            proc = self.run_script(
                [
                    PYTHON,
                    "scripts/material_passport.py",
                    "--artifact",
                    str(Path(tmp) / "future.md"),
                    "--scope",
                    "short",
                    "--output",
                    str(out),
                    "--no-event",
                ]
            )
            self.assertEqual(proc.returncode, 0, proc.stderr + proc.stdout)
            text = out.read_text(encoding="utf-8")
            self.assertIn("Status: `WARN`", text)
            self.assertIn("artifact: file does not exist yet", text)

    def test_full_material_passport_holds_for_missing_delivery_evidence(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            artifact = tmp_path / "draft.md"
            artifact.write_text("Clean draft.", encoding="utf-8")
            out = tmp_path / "passport.md"
            proc = self.run_script(
                [
                    PYTHON,
                    "scripts/material_passport.py",
                    "--artifact",
                    str(artifact),
                    "--scope",
                    "full",
                    "--output",
                    str(out),
                    "--no-event",
                ]
            )
            self.assertEqual(proc.returncode, 1)
            text = out.read_text(encoding="utf-8")
            self.assertIn("Status: `HOLD`", text)
            self.assertIn("runtime_receipt", text)

    def test_pre_delivery_lock_and_guard_pass_for_complete_basic_artifact(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            artifact = tmp_path / "output.docx"
            source = tmp_path / "source.md"
            self.write_basic_docx(artifact)
            source.write_text("A clean source paragraph without citations.", encoding="utf-8")
            evidence_files = {}
            for name in ["runtime", "passport", "source_map", "integrity", "quality"]:
                path = tmp_path / f"{name}.md"
                path.write_text(f"{name} evidence", encoding="utf-8")
                evidence_files[name] = path

            lock_proc = self.run_script(
                [
                    PYTHON,
                    "scripts/pre_delivery_lock.py",
                    "create",
                    "--target",
                    str(artifact),
                    "--runtime-receipt",
                    str(evidence_files["runtime"]),
                    "--material-passport",
                    str(evidence_files["passport"]),
                    "--source-map",
                    str(evidence_files["source_map"]),
                    "--integrity-preflight",
                    str(evidence_files["integrity"]),
                    "--quality-gate",
                    str(evidence_files["quality"]),
                    "--no-event",
                ]
            )
            self.assertEqual(lock_proc.returncode, 0, lock_proc.stderr + lock_proc.stdout)

            guard_proc = self.run_script(
                [
                    PYTHON,
                    "scripts/formal_delivery_guard.py",
                    "--artifact",
                    str(artifact),
                    "--source",
                    str(source),
                ]
            )
            self.assertEqual(guard_proc.returncode, 0, guard_proc.stderr + guard_proc.stdout)
            self.assertIn("pre_delivery_lock: PASS", guard_proc.stdout)

    def test_delivery_guard_blocks_without_lock(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            artifact = Path(tmp) / "output.docx"
            artifact.write_text("placeholder binary stand-in", encoding="utf-8")
            proc = self.run_script([PYTHON, "scripts/formal_delivery_guard.py", "--artifact", str(artifact)])
            self.assertEqual(proc.returncode, 1)
            self.assertIn("pre_delivery_lock: BLOCK", proc.stdout)

    def test_delivery_guard_blocks_when_required_skill_receipts_are_missing(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            artifact = tmp_path / "output.docx"
            source = tmp_path / "source.md"
            self.write_basic_docx(artifact)
            source.write_text("A clean source paragraph without citations.", encoding="utf-8")
            evidence_files = {}
            for name in ["runtime", "passport", "source_map", "integrity", "quality"]:
                path = tmp_path / f"{name}.md"
                path.write_text(f"{name} evidence", encoding="utf-8")
                evidence_files[name] = path

            lock_proc = self.run_script(
                [
                    PYTHON,
                    "scripts/pre_delivery_lock.py",
                    "create",
                    "--target",
                    str(artifact),
                    "--runtime-receipt",
                    str(evidence_files["runtime"]),
                    "--material-passport",
                    str(evidence_files["passport"]),
                    "--source-map",
                    str(evidence_files["source_map"]),
                    "--integrity-preflight",
                    str(evidence_files["integrity"]),
                    "--quality-gate",
                    str(evidence_files["quality"]),
                    "--no-event",
                ]
            )
            self.assertEqual(lock_proc.returncode, 0, lock_proc.stderr + lock_proc.stdout)

            task_id = f"missing-receipts-{uuid.uuid4().hex}"
            proc = self.run_script(
                [
                    PYTHON,
                    "scripts/formal_delivery_guard.py",
                    "--artifact",
                    str(artifact),
                    "--source",
                    str(source),
                    "--require-skill-receipts",
                    "--task-id",
                    task_id,
                ]
            )
            self.assertEqual(proc.returncode, 1)
            self.assertIn("skill_execution_receipts: BLOCK", proc.stdout)

    def test_delivery_guard_passes_when_required_skill_receipts_exist(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            artifact = tmp_path / "output.docx"
            source = tmp_path / "source.md"
            self.write_basic_docx(artifact)
            source.write_text("A clean source paragraph without citations.", encoding="utf-8")
            evidence_files = {}
            for name in ["runtime", "passport", "source_map", "integrity", "quality", "gate"]:
                path = tmp_path / f"{name}.md"
                path.write_text(f"{name} evidence", encoding="utf-8")
                evidence_files[name] = path

            lock_proc = self.run_script(
                [
                    PYTHON,
                    "scripts/pre_delivery_lock.py",
                    "create",
                    "--target",
                    str(artifact),
                    "--runtime-receipt",
                    str(evidence_files["runtime"]),
                    "--material-passport",
                    str(evidence_files["passport"]),
                    "--source-map",
                    str(evidence_files["source_map"]),
                    "--integrity-preflight",
                    str(evidence_files["integrity"]),
                    "--quality-gate",
                    str(evidence_files["quality"]),
                    "--no-event",
                ]
            )
            self.assertEqual(lock_proc.returncode, 0, lock_proc.stderr + lock_proc.stdout)

            task_id = f"complete-receipts-{uuid.uuid4().hex}"
            for receipt in [
                "dissertation-source-first-gate@thinking",
                "material-passport@thinking",
                "academic-integrity-preflight@thinking",
                "cognitive-frameworks@thinking",
                "academic-self-review-loop@writing",
                "authorial-voice-integrity@writing",
                "style-fingerprint-gate@writing",
                "uk-academic-writing-style@writing",
                "style-memory-and-revision-gate@writing",
                "dissertation-document-quality-gate@writing",
            ]:
                skill, stage = receipt.split("@", 1)
                receipt_proc = self.run_script(
                    [
                        PYTHON,
                        "scripts/skill_execution_receipt.py",
                        "create",
                        "--task-id",
                        task_id,
                        "--skill",
                        skill,
                        "--stage",
                        stage,
                        "--artifact",
                        str(artifact),
                        "--status",
                        "PASS",
                        "--evidence",
                        str(evidence_files["gate"]),
                        "--command",
                        "unit-test gate evidence",
                    ]
                )
                self.assertEqual(receipt_proc.returncode, 0, receipt_proc.stderr + receipt_proc.stdout)

            proc = self.run_script(
                [
                    PYTHON,
                    "scripts/formal_delivery_guard.py",
                    "--artifact",
                    str(artifact),
                    "--source",
                    str(source),
                    "--require-skill-receipts",
                    "--task-id",
                    task_id,
                ]
            )
            self.assertEqual(proc.returncode, 0, proc.stderr + proc.stdout)
            self.assertIn("skill_execution_receipts: PASS", proc.stdout)

    def test_delivery_guard_override_is_auditable_not_a_pass(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            artifact = Path(tmp) / "output.docx"
            artifact.write_text("placeholder binary stand-in", encoding="utf-8")
            proc = self.run_script(
                [
                    PYTHON,
                    "scripts/formal_delivery_guard.py",
                    "--artifact",
                    str(artifact),
                    "--acknowledge-override",
                    "--override-reason",
                    "Testing an explicit user-acknowledged exception.",
                ]
            )
            self.assertEqual(proc.returncode, 0, proc.stderr + proc.stdout)
            self.assertIn("pre_delivery_lock: BLOCK", proc.stdout)
            self.assertIn("delivery_override: ACKNOWLEDGED", proc.stdout)

    def test_delivery_guard_requires_reason_for_layout_exception_flags(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            artifact = tmp_path / "output.docx"
            source = tmp_path / "source.md"
            self.write_basic_docx(artifact)
            source.write_text("A clean source paragraph.", encoding="utf-8")

            for flag in [
                "--skip-structure-parity",
                "--skip-layout-review",
                "--allow-table-loss",
                "--allow-layout-regression",
            ]:
                with self.subTest(flag=flag):
                    proc = self.run_script(
                        [
                            PYTHON,
                            "scripts/formal_delivery_guard.py",
                            "--artifact",
                            str(artifact),
                            "--source",
                            str(source),
                            flag,
                        ]
                    )
                    self.assertEqual(proc.returncode, 2)
                    self.assertIn("--layout-decision-reason is required", proc.stdout)

    def test_delivery_guard_records_layout_exception_reason(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            artifact = tmp_path / "output.docx"
            source = tmp_path / "source.md"
            self.write_basic_docx(artifact)
            source.write_text("A clean source paragraph.", encoding="utf-8")

            proc = self.run_script(
                [
                    PYTHON,
                    "scripts/formal_delivery_guard.py",
                    "--artifact",
                    str(artifact),
                    "--source",
                    str(source),
                    "--skip-layout-review",
                    "--layout-decision-reason",
                    "Unit test confirms explicit logging for a skipped layout check.",
                ]
            )
            self.assertEqual(proc.returncode, 1)
            self.assertIn("pre_delivery_lock: BLOCK", proc.stdout)
            self.assertIn("layout_exception_reason: PASS", proc.stdout)

    def test_markdown_docx_structure_check_blocks_pipe_table_paragraphs(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source = tmp_path / "source.md"
            source.write_text("| A | B |\n|---|---|\n| one | two |\n", encoding="utf-8")
            artifact = tmp_path / "bad.docx"
            self.write_pipe_table_docx(artifact)

            proc = self.run_script(
                [
                    PYTHON,
                    "scripts/markdown_docx_structure_check.py",
                    "--markdown",
                    str(source),
                    "--docx",
                    str(artifact),
                    "--no-report",
                ]
            )
            self.assertEqual(proc.returncode, 1)
            self.assertIn("Status: BLOCK", proc.stdout)
            self.assertIn("pipe_table_paragraphs: 3", proc.stdout)

    def test_markdown_docx_structure_check_passes_real_word_table(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source = tmp_path / "source.md"
            source.write_text("| A | B |\n|---|---|\n| one | two |\n", encoding="utf-8")
            artifact = tmp_path / "good.docx"
            self.write_real_table_docx(artifact)

            proc = self.run_script(
                [
                    PYTHON,
                    "scripts/markdown_docx_structure_check.py",
                    "--markdown",
                    str(source),
                    "--docx",
                    str(artifact),
                    "--no-report",
                ]
            )
            self.assertEqual(proc.returncode, 0, proc.stderr + proc.stdout)
            self.assertIn("Status: PASS", proc.stdout)

    def test_docx_layout_review_blocks_heading_bold_override(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            artifact = Path(tmp) / "flat-heading.docx"
            self.write_flat_heading_docx(artifact)

            proc = self.run_script(
                [
                    PYTHON,
                    "scripts/docx_layout_review_check.py",
                    "--docx",
                    str(artifact),
                    "--no-report",
                ]
            )
            self.assertEqual(proc.returncode, 1)
            self.assertIn("Status: BLOCK", proc.stdout)
            self.assertIn("heading_issues: 1", proc.stdout)

    def test_formal_delivery_guard_blocks_missing_markdown_table_render(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            artifact = tmp_path / "output.docx"
            source = tmp_path / "source.md"
            source.write_text("| A | B |\n|---|---|\n| one | two |\n", encoding="utf-8")
            self.write_basic_docx(artifact)
            evidence_files = {}
            for name in ["runtime", "passport", "source_map", "integrity", "quality"]:
                path = tmp_path / f"{name}.md"
                path.write_text(f"{name} evidence", encoding="utf-8")
                evidence_files[name] = path

            lock_proc = self.run_script(
                [
                    PYTHON,
                    "scripts/pre_delivery_lock.py",
                    "create",
                    "--target",
                    str(artifact),
                    "--runtime-receipt",
                    str(evidence_files["runtime"]),
                    "--material-passport",
                    str(evidence_files["passport"]),
                    "--source-map",
                    str(evidence_files["source_map"]),
                    "--integrity-preflight",
                    str(evidence_files["integrity"]),
                    "--quality-gate",
                    str(evidence_files["quality"]),
                    "--no-event",
                ]
            )
            self.assertEqual(lock_proc.returncode, 0, lock_proc.stderr + lock_proc.stdout)

            proc = self.run_script(
                [
                    PYTHON,
                    "scripts/formal_delivery_guard.py",
                    "--artifact",
                    str(artifact),
                    "--source",
                    str(source),
                    "--skip-integrity-preflight",
                ]
            )
            self.assertEqual(proc.returncode, 1)
            self.assertIn("markdown_docx_structure_parity: BLOCK", proc.stdout)
            self.assertIn("docx_layout_review: BLOCK", proc.stdout)

    def test_pre_delivery_lock_check_detects_stale_evidence(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            artifact = tmp_path / "output.docx"
            self.write_basic_docx(artifact)
            evidence_files = {}
            for name in ["runtime", "passport", "source_map", "integrity", "quality"]:
                path = tmp_path / f"{name}.md"
                path.write_text(f"{name} evidence", encoding="utf-8")
                evidence_files[name] = path

            create_proc = self.run_script(
                [
                    PYTHON,
                    "scripts/pre_delivery_lock.py",
                    "create",
                    "--target",
                    str(artifact),
                    "--runtime-receipt",
                    str(evidence_files["runtime"]),
                    "--material-passport",
                    str(evidence_files["passport"]),
                    "--source-map",
                    str(evidence_files["source_map"]),
                    "--integrity-preflight",
                    str(evidence_files["integrity"]),
                    "--quality-gate",
                    str(evidence_files["quality"]),
                    "--no-event",
                ]
            )
            self.assertEqual(create_proc.returncode, 0, create_proc.stderr + create_proc.stdout)
            evidence_files["quality"].unlink()

            check_proc = self.run_script([PYTHON, "scripts/pre_delivery_lock.py", "check", "--target", str(artifact)])
            self.assertEqual(check_proc.returncode, 1)
            self.assertIn("quality_gate", check_proc.stdout)


if __name__ == "__main__":
    unittest.main()
