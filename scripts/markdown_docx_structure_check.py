#!/usr/bin/env python3
"""Check that Markdown structural features survive DOCX rendering."""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path

from _docx_runtime import ensure_python_docx

ensure_python_docx()
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "audit-reports" / "structure-parity"


@dataclass
class MarkdownTable:
    start_line: int
    rows: list[list[str]]


@dataclass
class MarkdownFeatures:
    tables: list[MarkdownTable]
    pipe_rows_outside_tables: int
    bullet_items: int
    numbered_items: int
    blockquote_blocks: int
    code_fence_blocks: int
    image_markers: int
    footnote_markers: int


@dataclass
class DocxFeatures:
    table_count: int
    table_cell_sets: list[set[str]]
    pipe_table_paragraphs: list[str]
    bullet_items: int
    numbered_items: int
    image_count: int


def rel(path: Path | None) -> str:
    if path is None:
        return "-"
    try:
        return str(path.resolve().relative_to(ROOT))
    except ValueError:
        return str(path.resolve())


def resolve(path_text: str | None) -> Path | None:
    if not path_text:
        return None
    path = Path(path_text)
    return path if path.is_absolute() else ROOT / path


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|") or not stripped.endswith("|"):
        return False
    inner = stripped.strip("|").replace(" ", "")
    return bool(inner) and all(char in "-:|" for char in inner)


def is_pipe_row(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|") or not stripped.endswith("|"):
        return False
    return stripped.count("|") >= 2


def split_table_row(line: str) -> list[str]:
    return [part.strip() for part in line.strip().strip("|").split("|")]


def strip_inline_markdown(text: str) -> str:
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.replace("**", "").replace("*", "").replace("`", "").strip()


def normalize_text(text: str) -> str:
    text = strip_inline_markdown(text)
    text = re.sub(r"\s+", " ", text)
    return text.strip().casefold()


def parse_markdown_features(path: Path) -> MarkdownFeatures:
    lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
    tables: list[MarkdownTable] = []
    pipe_rows_outside_tables = 0
    bullet_items = 0
    numbered_items = 0
    blockquote_blocks = 0
    code_fence_blocks = 0
    image_markers = 0
    footnote_markers = 0
    in_fence = False
    in_blockquote = False
    index = 0

    while index < len(lines):
        stripped = lines[index].strip()
        if stripped.startswith("```") or stripped.startswith("~~~"):
            if not in_fence:
                code_fence_blocks += 1
            in_fence = not in_fence
            index += 1
            continue
        if in_fence:
            index += 1
            continue

        if (
            is_pipe_row(stripped)
            and index + 1 < len(lines)
            and is_table_separator(lines[index + 1])
        ):
            start_line = index + 1
            rows: list[list[str]] = []
            while index < len(lines) and is_pipe_row(lines[index]):
                if not is_table_separator(lines[index]):
                    rows.append(split_table_row(lines[index]))
                index += 1
            tables.append(MarkdownTable(start_line=start_line, rows=rows))
            in_blockquote = False
            continue

        if is_pipe_row(stripped):
            pipe_rows_outside_tables += 1
        if re.match(r"^[-*+]\s+\S", stripped):
            bullet_items += 1
        if re.match(r"^\d+[.)]\s+\S", stripped):
            numbered_items += 1
        if stripped.startswith(">"):
            if not in_blockquote:
                blockquote_blocks += 1
            in_blockquote = True
        else:
            in_blockquote = False
        image_markers += len(re.findall(r"!\[[^\]]*\]\([^)]+\)", stripped))
        footnote_markers += len(re.findall(r"\[\^[^\]]+\]", stripped))
        index += 1

    return MarkdownFeatures(
        tables=tables,
        pipe_rows_outside_tables=pipe_rows_outside_tables,
        bullet_items=bullet_items,
        numbered_items=numbered_items,
        blockquote_blocks=blockquote_blocks,
        code_fence_blocks=code_fence_blocks,
        image_markers=image_markers,
        footnote_markers=footnote_markers,
    )


def table_signature(table: MarkdownTable) -> list[str]:
    if not table.rows:
        return []
    return [normalize_text(cell) for cell in table.rows[0] if normalize_text(cell)]


def paragraph_has_numbering(paragraph) -> bool:
    p_pr = paragraph._p.pPr
    return bool(p_pr is not None and p_pr.numPr is not None)


def parse_docx_features(path: Path) -> DocxFeatures:
    doc = Document(path)
    table_cell_sets: list[set[str]] = []
    for table in doc.tables:
        cells: set[str] = set()
        for row in table.rows:
            for cell in row.cells:
                normalized = normalize_text(cell.text)
                if normalized:
                    cells.add(normalized)
        table_cell_sets.append(cells)

    pipe_table_paragraphs: list[str] = []
    bullet_items = 0
    numbered_items = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if is_pipe_row(text):
            pipe_table_paragraphs.append(text)
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if "List Bullet" in style_name:
            bullet_items += 1
        if "List Number" in style_name or paragraph_has_numbering(paragraph):
            numbered_items += 1

    image_count = sum(1 for reln in doc.part.rels.values() if "image" in reln.reltype)
    return DocxFeatures(
        table_count=len(doc.tables),
        table_cell_sets=table_cell_sets,
        pipe_table_paragraphs=pipe_table_paragraphs,
        bullet_items=bullet_items,
        numbered_items=numbered_items,
        image_count=image_count,
    )


def signature_present(signature: list[str], docx: DocxFeatures) -> bool:
    if not signature:
        return True
    for cells in docx.table_cell_sets:
        if all(cell in cells for cell in signature):
            return True
    return False


def build_report(
    markdown_path: Path,
    docx_path: Path,
    previous_docx: Path | None,
    md: MarkdownFeatures,
    docx: DocxFeatures,
    previous: DocxFeatures | None,
    blockers: list[str],
    warnings: list[str],
) -> str:
    status = "BLOCK" if blockers else "PASS"
    table_rows = [
        ("tables", len(md.tables), docx.table_count, previous.table_count if previous else "-"),
        ("pipe table paragraphs", "-", len(docx.pipe_table_paragraphs), "-"),
        ("bullet items", md.bullet_items, docx.bullet_items, previous.bullet_items if previous else "-"),
        ("numbered items", md.numbered_items, docx.numbered_items, previous.numbered_items if previous else "-"),
        ("image markers", md.image_markers, docx.image_count, previous.image_count if previous else "-"),
    ]
    lines = [
        "# Markdown-DOCX Structural Parity Check",
        "",
        f"Generated: {datetime.now(timezone.utc).isoformat(timespec='seconds')}",
        f"Status: `{status}`",
        f"Markdown source: `{rel(markdown_path)}`",
        f"DOCX output: `{rel(docx_path)}`",
        f"Previous DOCX: `{rel(previous_docx)}`",
        "",
        "## Summary",
        "",
        "| Feature | Markdown | DOCX | Previous DOCX |",
        "|---|---:|---:|---:|",
    ]
    for feature, source_count, docx_count, previous_count in table_rows:
        lines.append(f"| {feature} | {source_count} | {docx_count} | {previous_count} |")
    lines.extend(["", "## Blockers", ""])
    lines.extend([f"- {item}" for item in blockers] if blockers else ["- None"])
    lines.extend(["", "## Warnings", ""])
    lines.extend([f"- {item}" for item in warnings] if warnings else ["- None"])
    lines.extend(
        [
            "",
            "## Boundary",
            "",
            "This check verifies structural rendering parity only. It does not verify citation claim support, academic quality, or submission readiness.",
            "",
        ]
    )
    return "\n".join(lines)


def main() -> int:
    parser = argparse.ArgumentParser(description="Check Markdown-to-DOCX structural parity.")
    parser.add_argument("--markdown", required=True, help="Source Markdown file.")
    parser.add_argument("--docx", required=True, help="Rendered DOCX file.")
    parser.add_argument("--previous-docx", help="Previous accepted DOCX for revision-safety comparison.")
    parser.add_argument("--allow-table-loss", action="store_true", help="Do not block when the previous DOCX had more tables than the current DOCX.")
    parser.add_argument("--output-dir", default=str(OUT_DIR), help="Directory for the parity report.")
    parser.add_argument("--no-report", action="store_true", help="Print result without writing a Markdown report.")
    args = parser.parse_args()

    markdown_path = resolve(args.markdown)
    docx_path = resolve(args.docx)
    previous_docx = resolve(args.previous_docx)
    if markdown_path is None or not markdown_path.exists():
        print(f"Missing Markdown source: {args.markdown}")
        return 2
    if docx_path is None or not docx_path.exists():
        print(f"Missing DOCX output: {args.docx}")
        return 2
    if args.previous_docx and (previous_docx is None or not previous_docx.exists()):
        print(f"Missing previous DOCX: {args.previous_docx}")
        return 2

    blockers: list[str] = []
    warnings: list[str] = []

    try:
        md = parse_markdown_features(markdown_path)
        docx = parse_docx_features(docx_path)
        previous = parse_docx_features(previous_docx) if previous_docx else None
    except Exception as exc:  # pragma: no cover - defensive CLI boundary
        print(f"Structural parity check failed to read inputs: {exc}")
        return 2

    if len(md.tables) > docx.table_count:
        blockers.append(f"Markdown has {len(md.tables)} table(s), but DOCX has {docx.table_count} real Word table(s).")
    for table in md.tables:
        signature = table_signature(table)
        if signature and not signature_present(signature, docx):
            blockers.append(f"Markdown table starting at line {table.start_line} is not matched by any DOCX table header.")
    if docx.pipe_table_paragraphs:
        blockers.append(f"DOCX contains {len(docx.pipe_table_paragraphs)} pipe-delimited table paragraph(s), indicating a table rendering failure.")
    if md.pipe_rows_outside_tables:
        warnings.append(f"Markdown contains {md.pipe_rows_outside_tables} pipe-delimited row(s) outside a recognised table.")
    if previous is not None and not args.allow_table_loss and previous.table_count > docx.table_count:
        blockers.append(f"Previous DOCX had {previous.table_count} table(s), but current DOCX has {docx.table_count}; pass --allow-table-loss only for an explicit layout decision.")
    if md.blockquote_blocks:
        warnings.append(f"Markdown contains {md.blockquote_blocks} blockquote block(s); visual render inspection should cover those pages.")
    if md.code_fence_blocks:
        warnings.append(f"Markdown contains {md.code_fence_blocks} code fence block(s); visual render inspection should cover those pages.")
    if md.footnote_markers:
        warnings.append(f"Markdown contains {md.footnote_markers} footnote marker(s); check that the DOCX rendering strategy is intentional.")

    report_path = None
    if not args.no_report:
        out_dir = resolve(args.output_dir) or OUT_DIR
        out_dir.mkdir(parents=True, exist_ok=True)
        report_path = out_dir / f"Markdown_DOCX_Structure_Check_{docx_path.stem}_{datetime.now().strftime('%Y-%m-%d_%H%M%S')}.md"
        report_path.write_text(
            build_report(markdown_path, docx_path, previous_docx, md, docx, previous, blockers, warnings),
            encoding="utf-8",
        )

    print(f"Status: {'BLOCK' if blockers else 'PASS'}")
    if report_path:
        print(f"Report: {report_path}")
    print(f"markdown_tables: {len(md.tables)}")
    print(f"docx_tables: {docx.table_count}")
    print(f"pipe_table_paragraphs: {len(docx.pipe_table_paragraphs)}")
    print(f"blockers: {len(blockers)}")
    print(f"warnings: {len(warnings)}")
    return 1 if blockers else 0


if __name__ == "__main__":
    raise SystemExit(main())
