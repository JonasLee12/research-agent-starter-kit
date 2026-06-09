#!/usr/bin/env python3
"""Check visible DOCX layout invariants before formal delivery.

This is a deterministic companion to the human visual render review. It catches
layout regressions that ordinary text or structural checks can miss, especially
heading run-level overrides that make a Word document look flatter than the
style definitions suggest.
"""

from __future__ import annotations

import argparse
import re
from collections import Counter
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path

from _docx_runtime import ensure_python_docx

ensure_python_docx()
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "audit-reports" / "layout-review"
EMU_PER_CM = 360000


@dataclass
class HeadingIssue:
    paragraph_index: int
    style_name: str
    text: str
    detail: str


@dataclass
class DocxLayout:
    heading_counts: Counter[str]
    table_count: int
    bullet_items: int
    numbered_items: int
    image_count: int
    body_paragraphs: int
    body_paragraphs_with_1_5_spacing: int
    page_width_cm: float | None
    page_height_cm: float | None
    margin_cm: tuple[float, float, float, float] | None
    heading_issues: list[HeadingIssue]


@dataclass
class MarkdownInventory:
    headings: Counter[str]
    tables: int
    bullet_items: int
    numbered_items: int
    code_fences: int
    blockquotes: int
    image_markers: int


def rel(path: Path | None) -> str:
    if path is None:
        return "-"
    try:
        return str(path.resolve().relative_to(ROOT))
    except ValueError:
        return str(path.resolve())


def resolve(path_text: str | None) -> Path | None:
    if not path_text:
        return None
    path = Path(path_text)
    return path if path.is_absolute() else ROOT / path


def cm(value) -> float | None:
    if value is None:
        return None
    return round(value / EMU_PER_CM, 2)


def is_heading_style(style_name: str) -> bool:
    return bool(re.match(r"^Heading\s+\d+$", style_name or ""))


def heading_level(style_name: str) -> str:
    match = re.match(r"^Heading\s+(\d+)$", style_name or "")
    return f"H{match.group(1)}" if match else style_name


def paragraph_has_numbering(paragraph) -> bool:
    p_pr = paragraph._p.pPr
    return bool(p_pr is not None and p_pr.numPr is not None)


def has_1_5_spacing(paragraph) -> bool:
    spacing = paragraph.paragraph_format.line_spacing
    if spacing is None and paragraph.style is not None:
        spacing = paragraph.style.paragraph_format.line_spacing
    if spacing is None:
        return False
    try:
        return abs(float(spacing) - 1.5) < 0.05
    except (TypeError, ValueError):
        return False


def inspect_docx(path: Path) -> DocxLayout:
    doc = Document(path)
    heading_counts: Counter[str] = Counter()
    heading_issues: list[HeadingIssue] = []
    bullet_items = 0
    numbered_items = 0
    body_paragraphs = 0
    body_paragraphs_with_1_5_spacing = 0

    for index, paragraph in enumerate(doc.paragraphs, start=1):
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if is_heading_style(style_name):
            level = heading_level(style_name)
            heading_counts[level] += 1
            style_bold = paragraph.style.font.bold
            if style_bold is False:
                heading_issues.append(
                    HeadingIssue(index, style_name, paragraph.text.strip(), "heading style is explicitly not bold")
                )
            if style_bold is None:
                heading_issues.append(
                    HeadingIssue(index, style_name, paragraph.text.strip(), "heading style has no explicit bold setting")
                )
            for run_index, run in enumerate(paragraph.runs, start=1):
                if paragraph.style.font.bold is True and run.font.bold is False:
                    heading_issues.append(
                        HeadingIssue(
                            index,
                            style_name,
                            paragraph.text.strip(),
                            f"run {run_index} explicitly sets bold=False and overrides the bold heading style",
                        )
                    )
            continue

        text = paragraph.text.strip()
        if text:
            body_paragraphs += 1
            if has_1_5_spacing(paragraph):
                body_paragraphs_with_1_5_spacing += 1
        if "List Bullet" in style_name:
            bullet_items += 1
        if "List Number" in style_name or paragraph_has_numbering(paragraph):
            numbered_items += 1

    section = doc.sections[0] if doc.sections else None
    page_width_cm = cm(section.page_width) if section else None
    page_height_cm = cm(section.page_height) if section else None
    margin_cm = None
    if section:
        margin_cm = (
            cm(section.top_margin) or 0.0,
            cm(section.bottom_margin) or 0.0,
            cm(section.left_margin) or 0.0,
            cm(section.right_margin) or 0.0,
        )
    image_count = sum(1 for reln in doc.part.rels.values() if "image" in reln.reltype)
    return DocxLayout(
        heading_counts=heading_counts,
        table_count=len(doc.tables),
        bullet_items=bullet_items,
        numbered_items=numbered_items,
        image_count=image_count,
        body_paragraphs=body_paragraphs,
        body_paragraphs_with_1_5_spacing=body_paragraphs_with_1_5_spacing,
        page_width_cm=page_width_cm,
        page_height_cm=page_height_cm,
        margin_cm=margin_cm,
        heading_issues=heading_issues,
    )


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|") or not stripped.endswith("|"):
        return False
    inner = stripped.strip("|").replace(" ", "")
    return bool(inner) and all(char in "-:|" for char in inner)


def is_pipe_row(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def parse_markdown(path: Path) -> MarkdownInventory:
    lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
    headings: Counter[str] = Counter()
    tables = 0
    bullet_items = 0
    numbered_items = 0
    code_fences = 0
    blockquotes = 0
    image_markers = 0
    in_fence = False
    in_blockquote = False
    index = 0

    while index < len(lines):
        stripped = lines[index].strip()
        if stripped.startswith("```") or stripped.startswith("~~~"):
            if not in_fence:
                code_fences += 1
            in_fence = not in_fence
            index += 1
            continue
        if in_fence:
            index += 1
            continue
        if is_pipe_row(stripped) and index + 1 < len(lines) and is_table_separator(lines[index + 1]):
            tables += 1
            while index < len(lines) and is_pipe_row(lines[index]):
                index += 1
            continue
        if match := re.match(r"^(#{1,6})\s+\S", stripped):
            headings[f"H{len(match.group(1))}"] += 1
        if re.match(r"^[-*+]\s+\S", stripped):
            bullet_items += 1
        if re.match(r"^\d+[.)]\s+\S", stripped):
            numbered_items += 1
        if stripped.startswith(">"):
            if not in_blockquote:
                blockquotes += 1
            in_blockquote = True
        else:
            in_blockquote = False
        image_markers += len(re.findall(r"!\[[^\]]*\]\([^)]+\)", stripped))
        index += 1

    return MarkdownInventory(
        headings=headings,
        tables=tables,
        bullet_items=bullet_items,
        numbered_items=numbered_items,
        code_fences=code_fences,
        blockquotes=blockquotes,
        image_markers=image_markers,
    )


def compare_counts(
    current: Counter[str],
    previous: Counter[str],
    label: str,
    blockers: list[str],
    warnings: list[str],
    allow_regression: bool,
) -> None:
    for key in sorted(set(current) | set(previous)):
        if previous[key] > current[key]:
            message = f"Previous DOCX had {previous[key]} {label} item(s) for {key}, current has {current[key]}."
            if allow_regression:
                warnings.append(f"Allowed layout regression: {message}")
            else:
                blockers.append(f"{message} Use --allow-layout-regression only for an explicit layout decision.")
        elif previous[key] != current[key]:
            warnings.append(f"{label} count changed for {key}: previous {previous[key]}, current {current[key]}.")


def build_report(
    docx_path: Path,
    markdown_path: Path | None,
    previous_docx: Path | None,
    current: DocxLayout,
    previous: DocxLayout | None,
    markdown: MarkdownInventory | None,
    blockers: list[str],
    warnings: list[str],
) -> str:
    status = "BLOCK" if blockers else "PASS"
    spacing_summary = "-"
    if current.body_paragraphs:
        spacing_summary = f"{current.body_paragraphs_with_1_5_spacing}/{current.body_paragraphs}"
    lines = [
        "# DOCX Layout Review Check",
        "",
        f"Generated: {datetime.now(timezone.utc).isoformat(timespec='seconds')}",
        f"Status: `{status}`",
        f"DOCX output: `{rel(docx_path)}`",
        f"Markdown source: `{rel(markdown_path)}`",
        f"Previous DOCX: `{rel(previous_docx)}`",
        "",
        "## Layout Inventory",
        "",
        "| Feature | Current DOCX | Previous DOCX | Markdown |",
        "|---|---:|---:|---:|",
        f"| Tables | {current.table_count} | {previous.table_count if previous else '-'} | {markdown.tables if markdown else '-'} |",
        f"| Bullet items | {current.bullet_items} | {previous.bullet_items if previous else '-'} | {markdown.bullet_items if markdown else '-'} |",
        f"| Numbered items | {current.numbered_items} | {previous.numbered_items if previous else '-'} | {markdown.numbered_items if markdown else '-'} |",
        f"| Images | {current.image_count} | {previous.image_count if previous else '-'} | {markdown.image_markers if markdown else '-'} |",
        f"| Body paragraphs with 1.5 spacing | {spacing_summary} | - | - |",
        "",
        "## Heading Counts",
        "",
        "| Level | Current DOCX | Previous DOCX | Markdown |",
        "|---|---:|---:|---:|",
    ]
    levels = sorted(
        set(current.heading_counts)
        | (set(previous.heading_counts) if previous else set())
        | (set(markdown.headings) if markdown else set())
    )
    for level in levels:
        lines.append(
            f"| {level} | {current.heading_counts[level]} | "
            f"{previous.heading_counts[level] if previous else '-'} | "
            f"{markdown.headings[level] if markdown else '-'} |"
        )
    lines.extend(
        [
            "",
            "## Page Setup",
            "",
            f"- Page size: `{current.page_width_cm} cm x {current.page_height_cm} cm`",
            f"- Margins top/bottom/left/right: `{current.margin_cm}`",
            "",
            "## Heading Issues",
            "",
        ]
    )
    if current.heading_issues:
        for issue in current.heading_issues:
            text = issue.text.replace("\n", " ")[:120]
            lines.append(
                f"- Paragraph {issue.paragraph_index}, {issue.style_name}: {issue.detail}; text: `{text}`"
            )
    else:
        lines.append("- None")
    lines.extend(["", "## Blockers", ""])
    lines.extend([f"- {item}" for item in blockers] if blockers else ["- None"])
    lines.extend(["", "## Warnings", ""])
    lines.extend([f"- {item}" for item in warnings] if warnings else ["- None"])
    lines.extend(
        [
            "",
            "## Boundary",
            "",
            "This deterministic check catches encodable layout regressions such as heading run overrides and structural count regressions. It does not replace visual inspection of rendered PDF/PNG pages or independent layout judgement when the layout self-review trigger fires.",
            "",
        ]
    )
    return "\n".join(lines)


def main() -> int:
    parser = argparse.ArgumentParser(description="Check DOCX layout invariants and baseline regressions.")
    parser.add_argument("--docx", required=True, help="Rendered DOCX file.")
    parser.add_argument("--markdown", help="Optional Markdown source.")
    parser.add_argument("--previous-docx", help="Previous accepted DOCX baseline.")
    parser.add_argument("--allow-layout-regression", action="store_true", help="Allow heading/table/list count regressions after an explicit logged layout decision.")
    parser.add_argument("--output-dir", default=str(OUT_DIR), help="Directory for the layout review report.")
    parser.add_argument("--no-report", action="store_true", help="Print result without writing a Markdown report.")
    args = parser.parse_args()

    docx_path = resolve(args.docx)
    markdown_path = resolve(args.markdown)
    previous_docx = resolve(args.previous_docx)
    if docx_path is None or not docx_path.exists():
        print(f"Missing DOCX output: {args.docx}")
        return 2
    if args.markdown and (markdown_path is None or not markdown_path.exists()):
        print(f"Missing Markdown source: {args.markdown}")
        return 2
    if args.previous_docx and (previous_docx is None or not previous_docx.exists()):
        print(f"Missing previous DOCX: {args.previous_docx}")
        return 2

    blockers: list[str] = []
    warnings: list[str] = []
    try:
        current = inspect_docx(docx_path)
        previous = inspect_docx(previous_docx) if previous_docx else None
        markdown = parse_markdown(markdown_path) if markdown_path else None
    except Exception as exc:  # pragma: no cover - defensive CLI boundary
        print(f"DOCX layout review failed to read inputs: {exc}")
        return 2

    for issue in current.heading_issues:
        if "bold=False" in issue.detail or "explicitly not bold" in issue.detail:
            blockers.append(f"Heading visual hierarchy issue at paragraph {issue.paragraph_index}: {issue.detail}.")
        else:
            warnings.append(f"Heading hierarchy warning at paragraph {issue.paragraph_index}: {issue.detail}.")

    if previous is not None:
        compare_counts(
            current.heading_counts,
            previous.heading_counts,
            "heading",
            blockers,
            warnings,
            args.allow_layout_regression,
        )
        if previous.table_count > current.table_count:
            message = f"Previous DOCX had {previous.table_count} table(s), current has {current.table_count}."
            if args.allow_layout_regression:
                warnings.append(f"Allowed layout regression: {message}")
            else:
                blockers.append(f"{message} Use --allow-layout-regression only for an explicit layout decision.")
        elif previous.table_count != current.table_count:
            warnings.append(f"Table count changed: previous {previous.table_count}, current {current.table_count}.")
        if previous.bullet_items > current.bullet_items:
            warnings.append(f"Bullet list count decreased: previous {previous.bullet_items}, current {current.bullet_items}.")
        if previous.numbered_items > current.numbered_items:
            warnings.append(f"Numbered list count decreased: previous {previous.numbered_items}, current {current.numbered_items}.")

    if markdown is not None:
        compare_counts(
            current.heading_counts,
            markdown.headings,
            "Markdown heading",
            blockers,
            warnings,
            args.allow_layout_regression,
        )
        if markdown.tables > current.table_count:
            blockers.append(f"Markdown has {markdown.tables} table(s), current DOCX has {current.table_count}.")
        if markdown.blockquotes:
            warnings.append(f"Markdown contains {markdown.blockquotes} blockquote block(s); rendered pages need visual inspection.")
        if markdown.code_fences:
            warnings.append(f"Markdown contains {markdown.code_fences} code fence block(s); rendered pages need visual inspection.")

    if current.body_paragraphs and current.body_paragraphs_with_1_5_spacing < current.body_paragraphs:
        warnings.append(
            f"{current.body_paragraphs - current.body_paragraphs_with_1_5_spacing} body paragraph(s) do not expose explicit 1.5 spacing through python-docx; verify rendered pages against formatting requirements."
        )

    report_path = None
    if not args.no_report:
        out_dir = resolve(args.output_dir) or OUT_DIR
        out_dir.mkdir(parents=True, exist_ok=True)
        report_path = out_dir / f"DOCX_Layout_Review_Check_{docx_path.stem}_{datetime.now().strftime('%Y-%m-%d_%H%M%S')}.md"
        report_path.write_text(
            build_report(docx_path, markdown_path, previous_docx, current, previous, markdown, blockers, warnings),
            encoding="utf-8",
        )

    print(f"Status: {'BLOCK' if blockers else 'PASS'}")
    if report_path:
        print(f"Report: {report_path}")
    print(f"heading_issues: {len(current.heading_issues)}")
    print(f"tables: {current.table_count}")
    print(f"blockers: {len(blockers)}")
    print(f"warnings: {len(warnings)}")
    return 1 if blockers else 0


if __name__ == "__main__":
    raise SystemExit(main())
